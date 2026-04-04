import os
import io
import json
import fitz  # PyMuPDF
import docx  # python-docx
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import Optional
from supabase import create_client, Client
from dotenv import load_dotenv
from openai import AsyncOpenAI

# Load environment variables
load_dotenv()

_supabase_url = os.getenv("SUPABASE_URL")
_supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
supabase: Optional[Client] = (
    create_client(_supabase_url, _supabase_key) if _supabase_url and _supabase_key else None
)
_groq_key = os.getenv("GROQ_API_KEY")

openai_client = AsyncOpenAI(
    api_key=_groq_key,
    base_url="https://api.groq.com/openai/v1"  # OpenAI-compatible Groq endpoint
)
router = APIRouter()

@router.post("/upload")
async def upload_resume(
    user_id: str = Form(...), 
    target_role: Optional[str] = Form(None), 
    file: UploadFile = File(...)
):
    """
    Scenario 1: Uploads a resume, extracts text, identifies skills.
    If target_role is missing, recommends jobs. Identifies gaps and preps data for the Plan router.
    """
    if not file.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF or DOCX files are allowed.")
    if not supabase:
        raise HTTPException(status_code=503, detail="Supabase is not configured on this server.")
    if not _groq_key:
        raise HTTPException(status_code=503, detail="GROQ_API_KEY is not configured on this server.")

    try:
        # 1. Read file contents into memory
        file_content = await file.read()
        extracted_text = ""
        
        # 2. Extract text based on file type
        if file.filename.lower().endswith('.pdf'):
            with fitz.open(stream=file_content, filetype="pdf") as pdf_doc:
                for page in pdf_doc:
                    extracted_text += page.get_text()
        elif file.filename.lower().endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_content))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
            
        extracted_text = " ".join(extracted_text.split())
        if not extracted_text:
            raise HTTPException(status_code=400, detail="Could not extract text. Ensure it is not an image-only PDF.")

        # 3. Dynamic GPT-4 Prompt formulation
        if target_role:
            system_instruction = f"Analyze this resume against the target role of '{target_role}'. Identify strengths and specifically highlight skill gaps preventing them from excelling in this role."
        else:
            system_instruction = "Analyze this resume. Identify the candidate's strengths, recommend 3 specific job roles they are highly suited for, and identify current skill gaps they need to fill to reach a senior level in the best-matching role."

        prompt = f"""
        {system_instruction}
        
        Resume Text:
        {extracted_text[:4000]}
        
        You MUST respond in strictly valid JSON format matching this schema exactly:
        {{
            "score": 85,
            "strengths": ["skill1", "skill2"],
            "target_role_evaluated": "The role evaluated or your top recommended role",
            "suggested_roles": ["role1", "role2", "role3"],
            "skill_gaps": ["gap1", "gap2"],
            "categories": [{{"subject": "Python", "A": 90, "fullMark": 100}}],
            "missingSkills": [{{"name": "Docker", "severity": "high"}}],
            "courses": [{{"title": "Docker for Beginners", "provider": "Coursera", "url": "https://coursera.org"}}]
        }}
        """

        # # 4. Call OpenAI GPT-4
        # completion = await openai_client.chat.completions.create(
        #     model="openai/gpt-oss-20b",
        #     messages=[
        #         {"role": "system", "content": "You are a precise, JSON-outputting career counselor AI."},
        #         {"role": "user", "content": prompt}
        #     ]
        # )
        
        # 4. Call the AI
        completion = await openai_client.chat.completions.create(
            # Make sure you are using a current Groq model like this one:
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a precise AI, career counselor AI. You MUST output ONLY raw JSON. No markdown, no formatting, no conversational text."},
                {"role": "user", "content": prompt}
            ],
            # If the model supports it, this forces JSON mode:
            response_format={"type": "json_object"} 
        )
        
        # 5. Parse the AI Response Defensively
        raw_content = completion.choices[0].message.content.strip()
        
        # Clean up Markdown code blocks if the AI added them
        if raw_content.startswith("```json"):
            raw_content = raw_content[7:] # Remove ```json
        elif raw_content.startswith("```"):
            raw_content = raw_content[3:] # Remove ```
            
        if raw_content.endswith("```"):
            raw_content = raw_content[:-3] # Remove ending ```
            
        raw_content = raw_content.strip() # Final trim of any whitespace
        
        try:
            ai_analysis = json.loads(raw_content)
            print(f"--- AI RAW OUTPUT ---\n{raw_content}\n---------------------")
        except json.JSONDecodeError:
            # Print the raw text to your terminal so you can see exactly how the AI messed up
            print(f"--- AI RAW OUTPUT ---\n{raw_content}\n---------------------")
            raise HTTPException(status_code=500, detail="AI failed to return valid JSON.")

        # 6. Save the result directly to Supabase
        db_response = supabase.table("resume_evaluations").insert({
            "user_id": user_id,
            "filename": file.filename,
            "analysis_result": ai_analysis
        }).execute()
        
        # 7. Return payload configured for frontend handoff to the Plan router
        return {
            "status": "success",
            "message": "Resume successfully processed. Ready for roadmap generation.",
            "data": db_response.data[0],
            "next_steps": {
                "action": "generate_plan",
                "endpoint": "/api/plan/generate",
                "payload_required": {
                    "user_id": user_id,
                    "target_role": ai_analysis.get("target_role_evaluated"),
                    "skill_gaps": ai_analysis.get("skill_gaps")
                }
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")


@router.get("/history/{user_id}")
async def get_resume_history(user_id: str):
    # ... (same as before) ...
    try:
        if not supabase:
            raise HTTPException(status_code=503, detail="Supabase is not configured on this server.")
        response = supabase.table("resume_evaluations").select("*").eq("user_id", user_id).execute()
        return {"status": "success", "user_id": user_id, "evaluations": response.data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching history: {str(e)}")


@router.post("/parse")
async def parse_resume(file: UploadFile = File(...)):
    """
    Parses a resume and returns structured data (skills, projects, experience, education, summary).
    This is a lightweight, stateless endpoint — nothing is saved to the database.
    """
    allowed_extensions = ('.pdf', '.docx', '.doc', '.txt')
    if not file.filename.lower().endswith(allowed_extensions):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, or TXT files are allowed.")
    if not _groq_key:
        raise HTTPException(status_code=503, detail="GROQ_API_KEY is not configured on this server.")

    try:
        # 1. Read file contents
        file_content = await file.read()
        extracted_text = ""

        # 2. Extract text based on file type
        if file.filename.lower().endswith('.pdf'):
            with fitz.open(stream=file_content, filetype="pdf") as pdf_doc:
                for page in pdf_doc:
                    extracted_text += page.get_text()
        elif file.filename.lower().endswith(('.docx', '.doc')):
            doc = docx.Document(io.BytesIO(file_content))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        elif file.filename.lower().endswith('.txt'):
            extracted_text = file_content.decode("utf-8", errors="ignore")

        extracted_text = " ".join(extracted_text.split())
        if not extracted_text:
            raise HTTPException(status_code=400, detail="Could not extract text from the file. Ensure it is not an image-only PDF.")

        # 3. AI Prompt for structured extraction
        prompt = f"""Extract structured information from the following resume text.

Resume Text:
{extracted_text[:5000]}

You MUST respond in strictly valid JSON matching this schema exactly:
{{
    "skills": ["skill1", "skill2", "skill3"],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Brief description of the project",
            "technologies": ["tech1", "tech2"]
        }}
    ],
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "duration": "Start - End",
            "description": "Brief description of responsibilities and achievements"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "University/College Name",
            "year": "Graduation Year or Duration"
        }}
    ],
    "summary": "A 2-3 sentence professional summary of the candidate based on the resume."
}}

Rules:
- Extract ALL skills mentioned (technical and soft skills).
- If a section is not found in the resume, return an empty array for that field.
- For summary, write a concise professional overview.
- Return ONLY raw JSON. No markdown, no extra text.
"""

        # 4. Call AI
        completion = await openai_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are an expert resume parser. You extract structured data from resumes and return ONLY valid JSON. No markdown, no conversational text."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )

        # 5. Parse AI response
        raw_content = completion.choices[0].message.content.strip()

        if raw_content.startswith("```json"):
            raw_content = raw_content[7:]
        elif raw_content.startswith("```"):
            raw_content = raw_content[3:]
        if raw_content.endswith("```"):
            raw_content = raw_content[:-3]
        raw_content = raw_content.strip()

        try:
            parsed_data = json.loads(raw_content)
            print(f"--- PARSE ENDPOINT RAW OUTPUT ---\n{raw_content}\n---------------------------------")
        except json.JSONDecodeError:
            print(f"--- PARSE ENDPOINT RAW OUTPUT (INVALID) ---\n{raw_content}\n---------------------------------")
            raise HTTPException(status_code=500, detail="AI failed to return valid JSON.")

        # Ensure all expected keys exist with defaults
        result = {
            "skills": parsed_data.get("skills", []),
            "projects": parsed_data.get("projects", []),
            "experience": parsed_data.get("experience", []),
            "education": parsed_data.get("education", []),
            "summary": parsed_data.get("summary", ""),
        }

        return result

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing resume: {str(e)}")